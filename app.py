import os
import time
import threading
from flask import Flask, request, render_template, send_file, Response, redirect, url_for
from docx import Document
import requests
import zipfile
from werkzeug.utils import secure_filename

app = Flask(__name__)

# フォルダ設定
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# 言語設定
LANGUAGE_CODES = {
    "日本語": "ja",
    "英語": "en",
    "中国語 (簡体字)": "zh-Hans",
    "中国語 (繁体字)": "zh-Hant",
    "ベトナム語": "vi",
    "フランス語": "fr",
    "ドイツ語": "de",
    "スペイン語": "es",
    "韓国語": "ko",
    "ロシア語": "ru",
    "イタリア語": "it",
    "ポルトガル語": "pt",
    "アラビア語": "ar",
}

# Microsoft Translator APIキー
TRANSLATOR_API_KEY = "1r2WZwFSKT5F2Qbvk40JSqVElgs8TP17pTMwfkWsNtS26KgkNPs0JQQJ99BAACi0881XJ3w3AAAbACOGHRrZ"
TRANSLATOR_REGION = "japaneast"

# 進捗管理
progress = {"current": 0, "total": 0, "paused": False}


def translate_word_file(input_file, output_file, source_language, target_language):
    """Wordファイルの翻訳処理"""
    doc = Document(input_file)

    # 段落と表内のテキストを収集
    paragraphs = [{"index": idx, "text": p.text.strip()} for idx, p in enumerate(doc.paragraphs) if p.text.strip()]
    table_texts = []
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_texts.append(cell.text.strip())
                    table_cells.append(cell)

    # 翻訳処理
    all_texts = [p["text"] for p in paragraphs] + table_texts
    translated_texts = translate_texts_with_retry(all_texts, source_language, target_language)

    # 翻訳結果を反映
    for p in paragraphs:
        doc.paragraphs[p["index"]].text = translated_texts.pop(0)
    for cell in table_cells:
        cell.text = translated_texts.pop(0)

    doc.save(output_file)


def translate_texts_with_retry(texts, source_language, target_language, max_retries=10, batch_size=50):
    """バッチ処理とリトライを組み合わせた翻訳処理"""
    translated_texts = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        for attempt in range(max_retries):
            try:
                translated_texts.extend(translate_texts(batch, source_language, target_language))
                break
            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 429:
                    time.sleep(2 ** attempt)  # リトライ時の待機時間
                else:
                    raise e
        else:
            raise Exception("最大再試行回数に達しました")
    return translated_texts


def translate_texts(texts, source_language, target_language):
    """Microsoft Translator APIを使用した翻訳"""
    endpoint = "https://api.cognitive.microsofttranslator.com/translate"
    headers = {
        "Ocp-Apim-Subscription-Key": TRANSLATOR_API_KEY,
        "Ocp-Apim-Subscription-Region": TRANSLATOR_REGION,
        "Content-type": "application/json",
    }
    params = {"api-version": "3.0", "from": source_language, "to": target_language}
    body = [{"text": text} for text in texts]

    response = requests.post(endpoint, headers=headers, params=params, json=body)
    response.raise_for_status()
    return [t["translations"][0]["text"] for t in response.json()]

@app.route("/", methods=["GET", "POST"])
def index():
    global progress
    if request.method == "POST":
        source_language = request.form.get("source_language")
        target_language = request.form.get("target_language")
        uploaded_files = request.files.getlist("files")

        if not source_language or not target_language or not uploaded_files:
            return render_template("index.html", languages=LANGUAGE_CODES.keys(), error="入力項目を確認してください。")

        # 初期化
        progress["current"] = 0
        progress["total"] = len(uploaded_files)
        progress["paused"] = False
        results = []

        def process_files():
            """ファイルを非同期で翻訳処理"""
            for file in uploaded_files:
                # ファイルの保存先
                input_path = os.path.join(UPLOAD_FOLDER, secure_filename(file.filename))
                output_path = os.path.join(OUTPUT_FOLDER, f"translated_{file.filename}")

                # ファイルを保存
                file.save(input_path)

                # 一時停止状態を確認
                while progress["paused"]:
                    time.sleep(0.5)  # 一時停止中はスリープ

                try:
                    # 翻訳処理
                    translate_word_file(
                        input_path, output_path, 
                        LANGUAGE_CODES[source_language], 
                        LANGUAGE_CODES[target_language]
                    )
                    results.append(os.path.basename(output_path))
                except Exception as e:
                    print(f"Error translating {file.filename}: {e}")
                    continue

                # 進捗更新
                progress["current"] += 1

        # スレッドで非同期実行
        threading.Thread(target=process_files).start()
        return render_template(
            "index.html", 
            languages=LANGUAGE_CODES.keys(), 
            files=results, 
            message="翻訳が進行中です。"
        )

    return render_template("index.html", languages=LANGUAGE_CODES.keys())



@app.route("/progress")
def progress_status():
    """進捗状況をクライアントに送信"""
    def generate():
        while progress["current"] < progress["total"]:
            yield f"data:{progress['current']}/{progress['total']}\n\n"
            time.sleep(0.5)  # 定期的に進捗を送信
        yield "data:complete\n\n"  # 完了通知
    return Response(generate(), content_type="text/event-stream")


@app.route("/pause", methods=["POST"])
def pause_translation():
    """翻訳処理を一時停止"""
    global progress
    progress["paused"] = True
    return "Paused"


@app.route("/resume", methods=["POST"])
def resume_translation():
    """翻訳処理を再開"""
    global progress
    progress["paused"] = False
    return "Resumed"


@app.route("/download_zip")
def download_zip():
    """翻訳されたファイルをZIP形式でダウンロード"""
    zip_path = os.path.join(OUTPUT_FOLDER, "translations.zip")
    with zipfile.ZipFile(zip_path, "w") as zipf:
        for file_name in os.listdir(OUTPUT_FOLDER):
            file_path = os.path.join(OUTPUT_FOLDER, file_name)
            zipf.write(file_path, os.path.basename(file_path))
    return send_file(zip_path, as_attachment=True)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
